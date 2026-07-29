from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src.schemas import WorkflowResult


@dataclass(frozen=True)
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    items: tuple[str, ...] = ()
    rows: tuple[tuple[str, ...], ...] = ()


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_UNORDERED_RE = re.compile(r"^\s*[-*+]\s+(.+?)\s*$")
_ORDERED_RE = re.compile(r"^\s*(\d+)[.)]\s+(.+?)\s*$")
_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
_INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|`.+?`|(?<!\*)\*[^*]+?\*(?!\*)|(?<!_)_[^_]+?_(?!_))")


def _normalise_markdown(markdown: str) -> str:
    text = markdown.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("—", "-").replace("–", "-")
    return text.strip()


def _split_table_row(line: str) -> tuple[str, ...]:
    stripped = line.strip().strip("|")
    return tuple(cell.strip() for cell in stripped.split("|"))


def parse_markdown(markdown: str) -> list[MarkdownBlock]:
    """Parse the classroom-focused Markdown subset used by EduAgent LK."""
    lines = _normalise_markdown(markdown).split("\n")
    blocks: list[MarkdownBlock] = []
    paragraph_lines: list[str] = []
    list_items: list[str] = []
    list_kind: str | None = None
    in_code_fence = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(MarkdownBlock("paragraph", " ".join(part.strip() for part in paragraph_lines)))
            paragraph_lines.clear()

    def flush_list() -> None:
        nonlocal list_kind
        if list_items and list_kind:
            blocks.append(MarkdownBlock(list_kind, items=tuple(list_items)))
            list_items.clear()
        list_kind = None

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            flush_list()
            if in_code_fence:
                blocks.append(MarkdownBlock("code", "\n".join(code_lines)))
                code_lines.clear()
                in_code_fence = False
            else:
                in_code_fence = True
            index += 1
            continue

        if in_code_fence:
            code_lines.append(line)
            index += 1
            continue

        heading = _HEADING_RE.match(line)
        unordered = _UNORDERED_RE.match(line)
        ordered = _ORDERED_RE.match(line)

        if heading:
            flush_paragraph()
            flush_list()
            blocks.append(MarkdownBlock("heading", heading.group(2), level=len(heading.group(1))))
            index += 1
            continue

        if "|" in line and index + 1 < len(lines) and _TABLE_SEPARATOR_RE.match(lines[index + 1]):
            flush_paragraph()
            flush_list()
            rows = [_split_table_row(line)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(_split_table_row(lines[index]))
                index += 1
            max_columns = max(len(row) for row in rows)
            padded = tuple(tuple(list(row) + [""] * (max_columns - len(row))) for row in rows)
            blocks.append(MarkdownBlock("table", rows=padded))
            continue

        if unordered:
            flush_paragraph()
            if list_kind not in (None, "unordered_list"):
                flush_list()
            list_kind = "unordered_list"
            list_items.append(unordered.group(1))
            index += 1
            continue

        if ordered:
            flush_paragraph()
            if list_kind not in (None, "ordered_list"):
                flush_list()
            list_kind = "ordered_list"
            list_items.append(ordered.group(2))
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            flush_list()
            index += 1
            continue

        if re.fullmatch(r"[-*_]{3,}", stripped):
            flush_paragraph()
            flush_list()
            blocks.append(MarkdownBlock("rule"))
            index += 1
            continue

        if stripped == "\\pagebreak":
            flush_paragraph()
            flush_list()
            blocks.append(MarkdownBlock("pagebreak"))
            index += 1
            continue

        flush_list()
        paragraph_lines.append(stripped)
        index += 1

    if in_code_fence and code_lines:
        blocks.append(MarkdownBlock("code", "\n".join(code_lines)))
    flush_paragraph()
    flush_list()
    return blocks


def _strip_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*|__(.+?)__", lambda m: m.group(1) or m.group(2), text)
    text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)|(?<!_)_([^_]+?)_(?!_)", lambda m: m.group(1) or m.group(2), text)
    return text.replace("`", "")


def _iter_inline_tokens(text: str) -> Iterable[tuple[str, str]]:
    cursor = 0
    for match in _INLINE_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            yield "normal", text[cursor:match.start()]
        token = match.group(0)
        if token.startswith(("**", "__")):
            yield "bold", token[2:-2]
        elif token.startswith("`"):
            yield "code", token[1:-1]
        else:
            yield "italic", token[1:-1]
        cursor = match.end()
    if cursor < len(text):
        yield "normal", text[cursor:]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def _add_docx_inline(paragraph, text: str) -> None:
    for token_type, value in _iter_inline_tokens(text):
        run = paragraph.add_run(_strip_markdown(value))
        if token_type == "bold":
            run.bold = True
        elif token_type == "italic":
            run.italic = True
        elif token_type == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(55, 65, 81)


def _configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    heading_sizes = {1: 20, 2: 15, 3: 12.5, 4: 11}
    for level, size in heading_sizes.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(20, 71, 122)
        style.paragraph_format.space_before = Pt(11 if level > 1 else 6)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True


def _add_docx_header(document: Document, result: WorkflowResult) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(result.draft.title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(21)
    title_run.font.color.rgb = RGBColor(20, 71, 122)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run("Curriculum-grounded material generated by EduAgent LK")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(9.5)
    subtitle_run.font.color.rgb = RGBColor(90, 101, 116)

    metadata = [
        ("Grade", str(result.request.grade)),
        ("Resource", result.request.output_type.value),
        ("Duration", f"{result.request.duration_minutes} minutes"),
        ("Student level", result.request.student_level.value),
        ("Quality score", f"{result.review.average_score:.1f} / 5"),
        ("Revision", str(result.draft.revision_number)),
    ]
    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, (label, value) in enumerate(metadata):
        cell = table.cell(index // 3, index % 3)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, "EEF5FB")
        _set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(f"{label}\n")
        label_run.bold = True
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(20, 71, 122)
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(9.5)
        value_run.font.color.rgb = RGBColor(31, 41, 55)

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_docx_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("EduAgent LK | Grade 9-10 English Teaching Assistant")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(110, 118, 129)


def _add_docx_markdown(document: Document, markdown: str) -> None:
    blocks = parse_markdown(markdown)
    first_heading_skipped = False
    for block in blocks:
        if block.kind == "heading":
            # The document already has a polished title, so avoid duplicating the top H1.
            if block.level == 1 and not first_heading_skipped:
                first_heading_skipped = True
                continue
            level = min(max(block.level - 1 if block.level > 1 else 1, 1), 4)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_docx_inline(paragraph, block.text)
        elif block.kind == "paragraph":
            paragraph = document.add_paragraph()
            _add_docx_inline(paragraph, block.text)
        elif block.kind in {"unordered_list", "ordered_list"}:
            for item_index, item in enumerate(block.items, start=1):
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.24)
                paragraph.paragraph_format.first_line_indent = Inches(-0.18)
                paragraph.paragraph_format.space_after = Pt(2)
                marker = f"{item_index}. " if block.kind == "ordered_list" else "• "
                marker_run = paragraph.add_run(marker)
                marker_run.bold = block.kind == "ordered_list"
                _add_docx_inline(paragraph, item)
        elif block.kind == "table" and block.rows:
            table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row_index, row in enumerate(block.rows):
                for column_index, value in enumerate(row):
                    cell = table.cell(row_index, column_index)
                    _set_cell_margins(cell, top=80, bottom=80)
                    if row_index == 0:
                        _set_cell_shading(cell, "DCEAF6")
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(0)
                    _add_docx_inline(paragraph, value)
                    if row_index == 0:
                        for run in paragraph.runs:
                            run.bold = True
            document.add_paragraph().paragraph_format.space_after = Pt(0)
        elif block.kind == "code":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            run = paragraph.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(55, 65, 81)
        elif block.kind == "rule":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run("_" * 75).font.color.rgb = RGBColor(190, 198, 208)
        elif block.kind == "pagebreak":
            document.add_page_break()


def build_docx_bytes(result: WorkflowResult) -> bytes:
    """Return a classroom-ready DOCX file as bytes for Streamlit download."""
    document = Document()
    _configure_docx_styles(document)
    _add_docx_header(document, result)
    _add_docx_markdown(document, result.draft.markdown)
    _add_docx_footer(document)

    core_properties = document.core_properties
    core_properties.title = result.draft.title
    core_properties.subject = f"Grade {result.request.grade} {result.request.output_type.value}"
    core_properties.author = "EduAgent LK"
    core_properties.keywords = "English teaching, Sri Lanka, Grade 9, Grade 10, agentic AI"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_inline(text: str) -> str:
    pieces: list[str] = []
    for token_type, value in _iter_inline_tokens(text):
        clean = escape(_strip_markdown(value))
        if token_type == "bold":
            pieces.append(f"<b>{clean}</b>")
        elif token_type == "italic":
            pieces.append(f"<i>{clean}</i>")
        elif token_type == "code":
            pieces.append(f'<font name="Courier" size="8.5">{clean}</font>')
        else:
            pieces.append(clean)
    return "".join(pieces)


def _pdf_styles():
    sample = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "EduTitle",
            parent=sample["Title"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#14477A"),
            alignment=TA_CENTER,
            spaceAfter=3 * mm,
        ),
        "subtitle": ParagraphStyle(
            "EduSubtitle",
            parent=sample["Normal"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#5A6574"),
            alignment=TA_CENTER,
            spaceAfter=4 * mm,
        ),
        "h1": ParagraphStyle(
            "EduH1",
            parent=sample["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14.5,
            leading=18,
            textColor=colors.HexColor("#14477A"),
            spaceBefore=4 * mm,
            spaceAfter=2 * mm,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "EduH2",
            parent=sample["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1E5C91"),
            spaceBefore=3.5 * mm,
            spaceAfter=1.7 * mm,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "EduH3",
            parent=sample["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=colors.HexColor("#2B5E86"),
            spaceBefore=2.5 * mm,
            spaceAfter=1.2 * mm,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "EduBody",
            parent=sample["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13.2,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=1.8 * mm,
            allowWidows=0,
            allowOrphans=0,
        ),
        "list": ParagraphStyle(
            "EduList",
            parent=sample["BodyText"],
            fontName="Helvetica",
            fontSize=9.3,
            leading=12.5,
            textColor=colors.HexColor("#1F2937"),
            leftIndent=3 * mm,
            spaceAfter=0.8 * mm,
        ),
        "code": ParagraphStyle(
            "EduCode",
            parent=sample["Code"],
            fontName="Courier",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#374151"),
            backColor=colors.HexColor("#F3F6F9"),
            borderPadding=5,
            leftIndent=4 * mm,
            rightIndent=4 * mm,
            spaceAfter=2 * mm,
        ),
        "table": ParagraphStyle(
            "EduTable",
            parent=sample["BodyText"],
            fontName="Helvetica",
            fontSize=8.2,
            leading=10.5,
            textColor=colors.HexColor("#1F2937"),
        ),
        "table_header": ParagraphStyle(
            "EduTableHeader",
            parent=sample["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.2,
            leading=10.5,
            textColor=colors.HexColor("#14477A"),
        ),
        "footer": ParagraphStyle(
            "EduFooter",
            parent=sample["Normal"],
            fontName="Helvetica",
            fontSize=7.5,
            textColor=colors.HexColor("#6E7681"),
            alignment=TA_CENTER,
        ),
    }


def _page_footer(canvas, document) -> None:
    canvas.saveState()
    width, _ = A4
    canvas.setStrokeColor(colors.HexColor("#D8E1EA"))
    canvas.setLineWidth(0.4)
    canvas.line(18 * mm, 13 * mm, width - 18 * mm, 13 * mm)
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#6E7681"))
    canvas.drawString(18 * mm, 8.5 * mm, "EduAgent LK | Grade 9-10 English Teaching Assistant")
    canvas.drawRightString(width - 18 * mm, 8.5 * mm, f"Page {document.page}")
    canvas.restoreState()


def _metadata_pdf_table(result: WorkflowResult, styles) -> Table:
    data = [
        [
            Paragraph("<b>Grade</b><br/>" + str(result.request.grade), styles["table"]),
            Paragraph("<b>Resource</b><br/>" + escape(result.request.output_type.value), styles["table"]),
            Paragraph("<b>Duration</b><br/>" + f"{result.request.duration_minutes} minutes", styles["table"]),
        ],
        [
            Paragraph("<b>Student level</b><br/>" + escape(result.request.student_level.value), styles["table"]),
            Paragraph("<b>Quality score</b><br/>" + f"{result.review.average_score:.1f} / 5", styles["table"]),
            Paragraph("<b>Revision</b><br/>" + str(result.draft.revision_number), styles["table"]),
        ],
    ]
    table = Table(data, colWidths=[56 * mm, 56 * mm, 56 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EEF5FB")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#C9DCEB")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8E4EE")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def _pdf_story(result: WorkflowResult, styles) -> list:
    story: list = [
        Paragraph(escape(result.draft.title), styles["title"]),
        Paragraph("Curriculum-grounded material generated by EduAgent LK", styles["subtitle"]),
        _metadata_pdf_table(result, styles),
        Spacer(1, 4 * mm),
    ]

    first_heading_skipped = False
    for block in parse_markdown(result.draft.markdown):
        if block.kind == "heading":
            if block.level == 1 and not first_heading_skipped:
                first_heading_skipped = True
                continue
            style_key = "h1" if block.level <= 2 else "h2" if block.level == 3 else "h3"
            story.append(Paragraph(_pdf_inline(block.text), styles[style_key]))
        elif block.kind == "paragraph":
            story.append(Paragraph(_pdf_inline(block.text), styles["body"]))
        elif block.kind in {"unordered_list", "ordered_list"}:
            for item_index, item in enumerate(block.items, start=1):
                marker = f"{item_index}." if block.kind == "ordered_list" else "•"
                story.append(Paragraph(_pdf_inline(item), styles["list"], bulletText=marker))
            story.append(Spacer(1, 1.2 * mm))
        elif block.kind == "table" and block.rows:
            table_data = []
            for row_index, row in enumerate(block.rows):
                row_style = styles["table_header"] if row_index == 0 else styles["table"]
                table_data.append([Paragraph(_pdf_inline(value), row_style) for value in row])
            available_width = A4[0] - 36 * mm
            column_width = available_width / len(block.rows[0])
            table = Table(table_data, colWidths=[column_width] * len(block.rows[0]), repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEAF6")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BFCEDB")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.extend([table, Spacer(1, 2 * mm)])
        elif block.kind == "code":
            story.append(Paragraph(escape(block.text).replace("\n", "<br/>"), styles["code"]))
        elif block.kind == "rule":
            rule = Table([[""]], colWidths=[A4[0] - 36 * mm], rowHeights=[0.5])
            rule.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#CAD6E1"))]))
            story.extend([Spacer(1, 1.5 * mm), rule, Spacer(1, 1.5 * mm)])
        elif block.kind == "pagebreak":
            story.append(PageBreak())
    return story


def build_pdf_bytes(result: WorkflowResult) -> bytes:
    """Return a polished A4 PDF as bytes without external system dependencies."""
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=15 * mm,
        bottomMargin=18 * mm,
        title=result.draft.title,
        author="EduAgent LK",
        subject=f"Grade {result.request.grade} {result.request.output_type.value}",
    )
    styles = _pdf_styles()
    document.build(
        _pdf_story(result, styles),
        onFirstPage=_page_footer,
        onLaterPages=_page_footer,
    )
    return output.getvalue()
