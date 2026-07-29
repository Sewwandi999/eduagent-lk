from __future__ import annotations

from io import BytesIO

from docx import Document

from src.exporters import build_docx_bytes, build_pdf_bytes, parse_markdown
from src.schemas import (
    AgentMessage,
    CurriculumContext,
    LessonDraft,
    OutputType,
    ReviewReport,
    RouteDecision,
    StudentLevel,
    TeacherRequest,
    WorkflowResult,
)


def sample_result() -> WorkflowResult:
    request = TeacherRequest(
        grade=9,
        topic="Reported Speech",
        duration_minutes=45,
        student_level=StudentLevel.AVERAGE,
        output_type=OutputType.LESSON_PLAN,
    )
    markdown = """# Grade 9 English: Reported Speech

## Learning Objectives
Students will be able to:
1. identify **reported speech**;
2. transform direct sentences accurately.

## Activity
- Work in pairs.
- Compare answers.

## Quick Table
| Direct | Reported |
|---|---|
| I am ready. | She said that she was ready. |

## Answer Key
1. She said that she was ready.
"""
    return WorkflowResult(
        request=request,
        route=RouteDecision(
            route="lesson_plan",
            requires_rag=True,
            complexity="standard",
            reason="test",
        ),
        curriculum=CurriculumContext(
            query="test",
            chunks=[],
            retrieval_notes="test",
            sufficient=True,
        ),
        draft=LessonDraft(
            title="Grade 9 Reported Speech - Lesson Plan",
            markdown=markdown,
            sections=["Learning Objectives", "Activity", "Quick Table", "Answer Key"],
            model_used="offline-template",
        ),
        review=ReviewReport(
            grade_suitability=5,
            grammar_accuracy=5,
            instruction_clarity=4,
            syllabus_alignment=4,
            answer_key_quality=5,
            strengths=["Clear"],
            issues=[],
            revision_required=False,
            model_used="offline-template",
        ),
        messages=[
            AgentMessage(
                sender="Orchestrator",
                receiver="TeacherUI",
                performative="result",
                task_id=request.task_id,
                payload={"ok": True},
            )
        ],
    )


def test_parse_markdown_supports_classroom_content():
    blocks = parse_markdown(sample_result().draft.markdown)
    kinds = [block.kind for block in blocks]
    assert "heading" in kinds
    assert "ordered_list" in kinds
    assert "unordered_list" in kinds
    assert "table" in kinds


def test_docx_export_is_valid_and_contains_content():
    payload = build_docx_bytes(sample_result())
    assert payload.startswith(b"PK")
    assert len(payload) > 10_000
    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Learning Objectives" in text
    assert "Answer Key" in text
    assert len(document.tables) >= 2


def test_pdf_export_is_valid():
    payload = build_pdf_bytes(sample_result())
    assert payload.startswith(b"%PDF")
    assert len(payload) > 2_000
